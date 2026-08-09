from __future__ import annotations

from pathlib import Path
from datetime import datetime
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[3]
REPORT_DIR = Path(__file__).resolve().parent
EVIDENCE = REPORT_DIR / "evidence"
ASSETS = REPORT_DIR / "assets"
OUTPUT_DOCX = ROOT / "output" / "product" / "SandPro_OMP_Pre_Rollout_Audit_2026-08-04.docx"
ASSETS.mkdir(parents=True, exist_ok=True)
OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

COLORS = {
    "ink": "080808",
    "muted": "5E6068",
    "soft": "F0F0F0",
    "white": "FFFFFF",
    "gold": "BBAA96",
    "lav": "CBA8FF",
    "lav_soft": "F4ECFF",
    "sand": "F6F1EA",
    "orange": "F58220",
    "orange_soft": "FFF1E5",
    "red": "B42318",
    "red_soft": "FDECEA",
    "amber": "9A6700",
    "amber_soft": "FFF4D6",
    "green": "167A5A",
    "green_soft": "E7F5EF",
    "blue": "3D5AA9",
    "blue_soft": "EBF0FF",
    "line": "D9D7D2",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["muted"])
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text, url, color="3D5AA9"):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_run(p, text, *, bold=False, color=None, size=None, italic=False, all_caps=False):
    r = p.add_run(text.upper() if all_caps else text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Arial"
    if color:
        r.font.color.rgb = rgb(color)
    if size:
        r.font.size = Pt(size)
    return r


def set_para(p, *, before=0, after=6, line=1.08, keep_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next


def add_label(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=2, after=3, keep_next=True)
    add_run(p, text, bold=True, color=COLORS["orange"], size=8, all_caps=True)
    return p


def add_heading(doc, text, level=1, subtitle=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    if subtitle:
        sp = doc.add_paragraph(subtitle, style="Subtitle")
        set_para(sp, after=10)
    return p


def add_body(doc, text, *, bold_lead=None, color=None, small=False):
    p = doc.add_paragraph()
    set_para(p, after=6, line=1.12)
    if bold_lead and text.startswith(bold_lead):
        add_run(p, bold_lead, bold=True, color=color or COLORS["ink"], size=8.8 if small else 9.6)
        add_run(p, text[len(bold_lead):], color=color or COLORS["ink"], size=8.8 if small else 9.6)
    else:
        add_run(p, text, color=color or COLORS["ink"], size=8.8 if small else 9.6)
    return p


def add_bullets(doc, items, *, color=None, tight=True):
    for item in items:
        p = doc.add_paragraph(style="NDAI Bullet")
        set_para(p, after=2.5 if tight else 5, line=1.08)
        add_run(p, item, color=color or COLORS["ink"], size=9)


def add_status_pill(cell, text, status):
    fill, fore = {
        "red": (COLORS["red_soft"], COLORS["red"]),
        "amber": (COLORS["amber_soft"], COLORS["amber"]),
        "green": (COLORS["green_soft"], COLORS["green"]),
        "blue": (COLORS["blue_soft"], COLORS["blue"]),
        "gray": (COLORS["soft"], COLORS["muted"]),
    }[status]
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=0)
    add_run(p, text, bold=True, color=fore, size=8, all_caps=True)


def add_metric_cards(doc, cards, columns=4):
    table = doc.add_table(rows=1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    set_repeat_table_header(table.rows[0])
    for i, (value, label, note, tone) in enumerate(cards):
        cell = table.rows[0].cells[i]
        shade(cell, {"red": COLORS["red_soft"], "amber": COLORS["amber_soft"], "green": COLORS["green_soft"], "lav": COLORS["lav_soft"], "sand": COLORS["sand"]}[tone])
        set_cell_border(cell, top={"val":"single","sz":"10","color":COLORS["white"]}, left={"val":"single","sz":"10","color":COLORS["white"]}, bottom={"val":"single","sz":"10","color":COLORS["white"]}, right={"val":"single","sz":"10","color":COLORS["white"]})
        set_cell_margins(cell, 120, 120, 120, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=1)
        add_run(p, value, bold=True, color=COLORS["ink"], size=18)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p2, after=2)
        add_run(p2, label, bold=True, color=COLORS["ink"], size=7.7, all_caps=True)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p3, after=0)
        add_run(p3, note, color=COLORS["muted"], size=7)
    return table


def add_callout(doc, title, body, tone="lav"):
    fills = {"lav": COLORS["lav_soft"], "red": COLORS["red_soft"], "amber": COLORS["amber_soft"], "green": COLORS["green_soft"], "sand": COLORS["sand"], "blue": COLORS["blue_soft"]}
    accents = {"lav": COLORS["lav"], "red": COLORS["red"], "amber": COLORS["amber"], "green": COLORS["green"], "sand": COLORS["gold"], "blue": COLORS["blue"]}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fills[tone])
    set_cell_border(cell, left={"val":"single","sz":"28","color":accents[tone]}, top={"val":"nil"}, right={"val":"nil"}, bottom={"val":"nil"})
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    set_para(p, after=3)
    add_run(p, title, bold=True, color=COLORS["ink"], size=10.5)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line=1.1)
    add_run(p2, body, color=COLORS["ink"], size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None, font_size=7.7, header_fill=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table)
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, header_fill or COLORS["ink"])
        set_cell_margins(cell, 90, 95, 90, 95)
        p = cell.paragraphs[0]
        set_para(p, after=0)
        add_run(p, h, bold=True, color=COLORS["white"], size=7.2, all_caps=True)
        if widths:
            cell.width = Inches(widths[i])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            shade(cell, COLORS["white"] if r_idx % 2 == 0 else "FAF9F7")
            set_cell_margins(cell, 80, 95, 80, 95)
            set_cell_border(cell, bottom={"val":"single","sz":"4","color":COLORS["line"]})
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.05)
            if isinstance(val, tuple):
                text, tone = val
                add_run(p, text, bold=True, color={"red":COLORS["red"],"amber":COLORS["amber"],"green":COLORS["green"],"blue":COLORS["blue"]}.get(tone,COLORS["ink"]), size=font_size)
            else:
                add_run(p, str(val), color=COLORS["ink"], size=font_size)
            if widths:
                cell.width = Inches(widths[i])
    return table


def add_figure(doc, path: Path, caption: str, width=6.95):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=2, after=2)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption[:120])
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(c, after=8)
    add_run(c, caption, italic=True, color=COLORS["muted"], size=7.5)


def new_page(doc, label, heading, subtitle=None):
    p = add_label(doc, label)
    p.paragraph_format.page_break_before = True
    add_heading(doc, heading, 1, subtitle)


def _font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/SFNS.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            pass
    return ImageFont.load_default()


def draw_bar_chart(out_path: Path, title: str, subtitle: str, labels, values, colors, max_value: int):
    width, height = 1600, 650
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 28), title, fill="#080808", font=_font(34, True))
    draw.text((40, 78), subtitle, fill="#5E6068", font=_font(20))
    left, right, top = 340, 1530, 145
    row_h = (height - top - 45) / len(labels)
    for tick in range(0, max_value + 1, max(1, max_value // 5)):
        x = left + (right - left) * tick / max_value
        draw.line((x, top, x, height - 30), fill="#ECEAE6", width=2)
        draw.text((x - 9, height - 28), str(tick), fill="#8A8A8A", font=_font(15))
    for i, (label, value, color) in enumerate(zip(labels, values, colors)):
        y = top + i * row_h + 14
        draw.text((40, y + 9), label, fill="#080808", font=_font(22, True))
        bar_w = max(3, int((right - left) * value / max_value))
        draw.rounded_rectangle((left, y, left + bar_w, y + 48), radius=10, fill=color)
        txt = str(value)
        draw.text((left + bar_w + 14, y + 9), txt, fill="#080808", font=_font(22, True))
    img.save(out_path, quality=92)


def make_charts():
    # Activation uses a single stacked bar because the disproportion is the message.
    img = Image.new("RGB", (1600, 430), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 30), "Activation is the dominant rollout constraint", fill="#080808", font=_font(34, True))
    draw.text((40, 82), "90 SandPro personnel accounts · 7 support / QA accounts excluded", fill="#5E6068", font=_font(20))
    x0, y0, total_w, bar_h = 40, 170, 1500, 92
    never_w = int(total_w * 85 / 90)
    draw.rounded_rectangle((x0, y0, x0 + total_w, y0 + bar_h), radius=18, fill="#CBA8FF")
    draw.rounded_rectangle((x0, y0, x0 + never_w, y0 + bar_h), radius=18, fill="#E9E4DE")
    draw.text((x0 + never_w / 2 - 115, y0 + 28), "85 never signed in", fill="#080808", font=_font(25, True))
    draw.text((x0 + never_w + 26, y0 + 28), "5", fill="#080808", font=_font(25, True))
    draw.text((40, 304), "Only 1 of 90 personnel signed in during the last seven days.", fill="#B42318", font=_font(24, True))
    img.save(ASSETS / "activation.png", quality=92)

    draw_bar_chart(ASSETS / "objectives.png", "Objective portfolio · 92 total", "82 active · 10 completed · 90 without due dates", ["On track", "At risk", "Not started", "Completed"], [55, 13, 14, 10], ["#167A5A", "#D99A2B", "#BBAA96", "#3D5AA9"], 60)
    draw_bar_chart(ASSETS / "ncr.png", "NCR closeout readiness", "296 closed · 74 open · all open rows need data", ["Total NCRs", "Open", "Critical open", "Past due", "Ownerless open", "Ready to close"], [370, 74, 48, 45, 73, 0], ["#D8D3CC", "#D99A2B", "#B42318", "#C86B3C", "#CBA8FF", "#167A5A"], 400)
    draw_bar_chart(ASSETS / "notifications.png", "Notification reach is opt-in constrained", "Channel-enabled SandPro personnel out of 90", ["In-app preference", "Email enabled", "Push enabled", "Active push subscription"], [90, 3, 3, 2], ["#3D5AA9", "#BBAA96", "#CBA8FF", "#F58220"], 100)


def make_montage(paths, labels, out_path, cols=2, max_width=1500, bg="#F0F0F0"):
    imgs = [Image.open(p).convert("RGB") for p in paths]
    thumb_w = (max_width - 60 - (cols - 1) * 24) // cols
    rows = (len(imgs) + cols - 1) // cols
    resized = []
    for im in imgs:
        ratio = thumb_w / im.width
        resized.append(im.resize((thumb_w, int(im.height * ratio)), Image.Resampling.LANCZOS))
    row_heights = []
    for r in range(rows):
        row_heights.append(max(x.height for x in resized[r*cols:(r+1)*cols]) + 42)
    canvas = Image.new("RGB", (max_width, 40 + sum(row_heights) + 20), bg)
    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 24)
    except Exception:
        font = ImageFont.load_default()
    y = 26
    for idx, im in enumerate(resized):
        r, c = divmod(idx, cols)
        if c == 0 and idx > 0:
            y += row_heights[r-1]
        x = 30 + c * (thumb_w + 24)
        draw.text((x, y), labels[idx], fill="#080808", font=font)
        canvas.paste(im, (x, y + 34))
    canvas.save(out_path, quality=90)


def build_assets():
    make_charts()
    make_montage(
        [EVIDENCE/"desktop/03-okr.jpg", EVIDENCE/"desktop/24-okr-settled.jpg", EVIDENCE/"desktop/23-kpi.jpg", EVIDENCE/"desktop/25-kpi-settled.jpg", EVIDENCE/"mobile/05-ncr.jpg", EVIDENCE/"mobile/07-ncr-settled.jpg"],
        ["OKR · false empty", "OKR · hydrated", "KPI · false zero", "KPI · hydrated", "Mobile NCR · 5.5 sec", "Mobile NCR · 7.4 sec"],
        ASSETS/"hydration-comparison.jpg", cols=2
    )
    make_montage(
        [EVIDENCE/"mobile/02-dashboard-settled.jpg", EVIDENCE/"mobile/03-create.jpg", EVIDENCE/"mobile/04-okr.jpg", EVIDENCE/"mobile/07-ncr-settled.jpg", EVIDENCE/"mobile/06-organization.jpg"],
        ["Dashboard", "Create", "OKR", "NCR", "Organization"],
        ASSETS/"mobile-surfaces.jpg", cols=3, max_width=1500
    )


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(.55)
    section.bottom_margin = Inches(.55)
    section.left_margin = Inches(.62)
    section.right_margin = Inches(.62)
    section.header_distance = Inches(.25)
    section.footer_distance = Inches(.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["ink"])
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(10)
    subtitle.font.color.rgb = rgb(COLORS["muted"])
    subtitle.paragraph_format.space_after = Pt(10)

    for level, size in ((1, 22), (2, 15), (3, 11)):
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = rgb(COLORS["ink"])
        s.paragraph_format.space_before = Pt(7 if level == 1 else 5)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    if "NDAI Bullet" not in [s.name for s in doc.styles]:
        bs = doc.styles.add_style("NDAI Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bs = doc.styles["NDAI Bullet"]
    bs.font.name = "Arial"
    bs.font.size = Pt(9)
    bs.paragraph_format.left_indent = Inches(.2)
    bs.paragraph_format.first_line_indent = Inches(-.12)
    bs.paragraph_format.space_after = Pt(2)
    num_pr = bs.element.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)

    # Header / footer: restrained NDAI light editorial system.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para(hp, after=0)
    add_run(hp, "ND.AI", bold=True, color=COLORS["ink"], size=8)
    add_run(hp, "   /   SANDPRO OMP   /   FINAL PRE-ROLLOUT AUDIT", color=COLORS["muted"], size=7.5)
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def cover(doc: Document):
    # Visual masthead
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(t)
    set_repeat_table_header(t.rows[0])
    left, right = t.rows[0].cells
    left.width = Inches(5.4)
    right.width = Inches(1.7)
    for c in (left, right):
        set_cell_margins(c, 110, 130, 110, 130)
        set_cell_border(c, top={"val":"nil"},left={"val":"nil"},right={"val":"nil"},bottom={"val":"single","sz":"18","color":COLORS["lav"]})
    p = left.paragraphs[0]
    set_para(p, after=0)
    add_run(p, "ND.AI", bold=True, color=COLORS["ink"], size=16)
    add_run(p, "   OPERATIONS AUDIT", bold=True, color=COLORS["muted"], size=7.5, all_caps=True)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(rp, after=0)
    add_run(rp, "CONFIDENTIAL", bold=True, color=COLORS["red"], size=7.5, all_caps=True)

    p = doc.add_paragraph()
    set_para(p, before=28, after=5)
    add_run(p, "FINAL PRE-ROLLOUT AUDIT", bold=True, color=COLORS["orange"], size=9, all_caps=True)
    p = doc.add_paragraph(style="Title")
    p.add_run("SandPro OMP\nproduction readiness")
    p = doc.add_paragraph()
    set_para(p, after=18)
    add_run(p, "objectivetracker.net  ·  observed production  ·  August 4, 2026", color=COLORS["muted"], size=11)

    add_callout(doc, "DECISION: HOLD the unguarded company-wide launch", "Proceed tomorrow only as a controlled, staffed rollout after the hard gates in this report are cleared or explicitly accepted by accountable owners. The application is visually strong and core surfaces function after hydration; adoption, first-load truth, communication freshness, and operational data stewardship are not yet company-wide ready.", "red")

    add_metric_cards(doc, [
        ("5 / 90", "Personnel ever signed in", "85 have never signed in", "red"),
        ("~7.4s", "False-zero NCR window", "Measured at mobile viewport", "amber"),
        ("74", "Open NCRs", "48 critical · 45 past due", "red"),
        ("Aug 1", "Current production deploy", "Vercel READY", "lav"),
    ])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    set_para(p, before=14, after=5)
    add_run(p, "WHAT IS READY", bold=True, color=COLORS["ink"], size=8, all_caps=True)
    add_bullets(doc, [
        "Core desktop and responsive mobile surfaces render cleanly and the Create flow opens without a write.",
        "Dashboard, OKR, KPI, NCR, Organization, Admin groups, account settings, and Fix-It all hydrate with production data.",
        "Production deployment is READY; no Vercel error or warning log entries were returned for the last 24 hours.",
        "Relevant public tables have RLS enabled; Fix-It has 0 active and 28 archived items.",
    ])
    p = doc.add_paragraph()
    set_para(p, before=10, after=4)
    add_run(p, "WHAT IS NOT READY", bold=True, color=COLORS["ink"], size=8, all_caps=True)
    add_bullets(doc, [
        "Company-wide authentication and onboarding: only 5 SandPro personnel have ever signed in.",
        "First-load truth: several routes display credible zero/empty states for 5–8 seconds before real data appears.",
        "Rollout communication: the auto-open Daily bulletin still announces Wednesday, June 24 as the company-wide launch.",
        "NCR operating ownership: 73 of 74 open NCRs have no owner and none are ready for closeout.",
        "Real-employee, ordinary-role, installed-PWA, inbox/junk, and device-push proof remains unobserved in this audit.",
    ])

    p = doc.add_paragraph()
    set_para(p, before=16, after=0)
    add_run(p, "Prepared for Andrew Emmel / NDAI  ·  read-only production audit  ·  evidence captured 08:30–09:02 CDT", color=COLORS["muted"], size=7.5)


def build_report():
    build_assets()
    doc = Document()
    configure_doc(doc)
    cover(doc)

    new_page(doc, "01 / DECISION", "The decision and the minimum safe path", "This is not a recommendation to cancel adoption. It is a recommendation to change the launch shape until the evidence matches the promise.")
    add_callout(doc, "Recommended launch posture", "NO-GO for an unguarded all-company launch in which 90 employees are expected to self-activate and operate independently. CONDITIONAL GO for a controlled, staffed, department-by-department activation wave after the seven hard gates below are green or formally risk-accepted.", "red")
    add_heading(doc, "Seven hard gates before the first broad cohort", 2)
    gate_rows = [
        (("RED", "red"), "Daily bulletin", "Replace or retire the June 24 rollout story; verify a fresh bulletin in production."),
        (("RED", "red"), "Activation", "Named owner, credential path, and successful sign-in proof for every first-wave user."),
        (("RED", "red"), "Ordinary-role proof", "Validate a real employee/member account across visibility, Create, notifications, and org surfaces."),
        (("RED", "red"), "First-load truth", "Use an explicit loading state or brief every facilitator never to treat early zeros as final data."),
        (("RED", "red"), "Device / channel proof", "One actual employee device: install/launch PWA, enable push, receive a push, receive email, inspect junk."),
        (("AMBER", "amber"), "Security acceptance", "Close or formally accept exposed SECURITY DEFINER execute grants and leaked-password protection risk."),
        (("AMBER", "amber"), "NCR operating posture", "Name the accountable triage owner and communicate that the historical backlog is not launch-day work."),
    ]
    add_table(doc, ["State", "Gate", "Acceptance evidence"], gate_rows, widths=[.7,1.35,4.9], font_size=8)
    add_heading(doc, "What would make tomorrow safe", 2)
    add_bullets(doc, [
        "Start with 5–10 employees, not all 90. Hold the next wave until every person signs in and completes the same short checklist.",
        "Keep one launch commander, one activation lead, one product/support lead, and one NCR/data steward in the room or on call.",
        "Treat login failure, wrong data visibility, false zeros lasting more than 15 seconds, or notification delivery failure as stop-the-line events.",
        "Pause the next cohort; do not attempt a data rollback. Preserve the working pilot group while the issue is diagnosed.",
    ], tight=False)

    new_page(doc, "02 / SCOPE", "What was audited — and what was deliberately not changed", "Evidence is limited to observed production, authoritative read-only data, and read-only deployment/platform inspection.")
    add_heading(doc, "Surfaces inspected", 2)
    add_table(doc, ["Surface", "Identity / viewport", "Observed"], [
        ("Entry + Daily", "Release Smoke Admin · desktop", "Auto-open bulletin, current date, stale rollout story"),
        ("Tasks & Projects", "Desktop + 390×844", "KPI cards, list, filters, ownerless-NCR banner"),
        ("Create", "Desktop + 390×844", "Wizard opened; no record submitted"),
        ("OKR", "Desktop + 390×844", "90-line August check-in view; four company OKRs visible"),
        ("KPI Command Center", "Desktop", "Computed objectives/NCR lens; manual scorecard empty"),
        ("NCR", "Desktop + 390×844", "Tracker, Closeout, Analytics, KPA Import, Dept triage"),
        ("Organization + Admin", "Desktop + 390×844", "97-person org, users, 10 assignment groups"),
        ("Notifications + account", "Desktop", "Empty personal feed; push enable setting; profile menu"),
        ("Fix-It", "Desktop", "0 active / 28 archived after hydration"),
    ], widths=[1.55,2.15,3.25], font_size=8)
    add_heading(doc, "Authoritative checks", 2)
    add_bullets(doc, [
        "Supabase project whgrkfhuzgwmbelocnhq: read-only SQL, table counts, auth cohort, notification coverage, NCR completeness, RLS presence, security and performance advisors.",
        "Vercel deployment inspection: production alias, deployment ID, READY state, build region/runtime, cron configuration, recent error/warning query.",
        "Live HTTP/PWA inspection: response headers, manifest, service worker, cache strategy, icons/shortcuts declarations.",
        "Repository provenance check: current HEAD and dirty-worktree state only; no build, test, migration, commit, deploy, or cleanup was run.",
    ])
    add_callout(doc, "Audit boundary", "No production records were edited, no forms were submitted, no imports or exports were triggered, no notifications were sent, no users were invited or changed, and no Fix-It items were archived. This was not a physical-device test and did not prove any employee mailbox, Microsoft 365 trace, or installed-PWA session.", "sand")

    new_page(doc, "03 / PRODUCTION TRUTH", "The current production snapshot", "All figures below are current observed values from August 4, 2026, not targets or presentation estimates.")
    add_metric_cards(doc, [
        ("97", "Auth users / profiles", "96 confirmed", "lav"),
        ("92", "Objectives", "82 active · 10 complete", "green"),
        ("370", "NCR records", "296 closed · 74 open", "amber"),
        ("1,810", "Notifications", "Historical records", "sand"),
    ])
    add_heading(doc, "Deployment and runtime", 2)
    add_table(doc, ["Item", "Observed production value", "Assessment"], [
        ("Primary URL", "https://objectivetracker.net", ("READY", "green")),
        ("Vercel deployment", "dpl_BxFsSyqhf7V5TsM89uHqBJWkAy8V", ("READY", "green")),
        ("Deployment created", "Aug 1, 2026 · 4:24:41 PM CDT", ("3 days old", "blue")),
        ("Runtime", "Vite static app + Node 24.x serverless functions", ("Observed", "blue")),
        ("Crons", "Daily digest 13:00 UTC weekdays; reminders 14:00 UTC daily", ("Configured", "green")),
        ("Local repository HEAD", "6b0bbe9 · keep OKR headers visible · Jul 29", ("Not proven live", "amber")),
        ("Deploy-to-commit linkage", "No Git commit metadata returned by Vercel inspection", ("Traceability gap", "amber")),
    ], widths=[1.5,3.9,1.55], font_size=8)
    add_heading(doc, "Production controls that are already healthy", 2)
    add_bullets(doc, [
        "RLS is enabled on relevant public data tables inspected in production.",
        "Vercel returned no production error or warning log entries in the last 24-hour query window.",
        "The sampled last 100 auth events contained 99 informational events and one transient 500 ('context canceled') during the audit.",
        "The Fix-It queue is operationally clean: 0 active, 28 archived.",
        "All 92 objectives have owners; none are marked blocked.",
        "Ten active assignment groups exist, including All Personnel and Office Personnel.",
    ])

    new_page(doc, "04 / ADOPTION", "Activation is the true company-wide blocker", "The accounts exist. The usage evidence does not.")
    add_figure(doc, ASSETS/"activation.png", "Production auth cohort. Support and QA identities are excluded from the 90-person SandPro population.", 6.9)
    add_table(doc, ["Cohort", "Total", "Ever signed in", "Never signed in", "Last 7 days"], [
        ("All auth users", "97", "11", "86", "5"),
        ("SandPro personnel", "90", "5", "85", "1"),
        ("Support / QA", "7", "6", "1", "4"),
    ], widths=[2.0,1.0,1.25,1.35,1.35], font_size=8.5)
    add_heading(doc, "Department exposure", 2)
    add_table(doc, ["Department", "Ever / total", "Rollout implication"], [
        ("Field", ("0 / 39", "red"), "Largest cohort has no prior sign-in proof."),
        ("Wellhead Shop", ("0 / 10", "red"), "No sign-in proof."),
        ("CP Warehouse", ("0 / 9", "red"), "No sign-in proof."),
        ("Flowback Shop", ("0 / 7", "red"), "No sign-in proof."),
        ("Automation", ("1 / 6", "amber"), "One prior user; five unproven."),
        ("Sales", ("1 / 5", "amber"), "One prior user; four unproven."),
        ("Admin", ("1 / 4", "amber"), "Admin coverage is thin."),
        ("Leadership", ("1 / 2", "amber"), "One of two proven."),
        ("Quality", ("1 / 3", "amber"), "One of three proven."),
    ], widths=[1.6,1.2,4.15], font_size=8)
    add_callout(doc, "Important interpretation", "97 accounts and 96 confirmed addresses do not equal 90 rollout-ready employees. A password hash proves account provisioning, not that the employee knows the credential, can reset it, can see the right data, or can complete work on their actual device.", "red")
    add_heading(doc, "Activation script for each wave", 2)
    add_bullets(doc, [
        "Sign in as the employee — not an admin or QA account.",
        "Open Tasks & Projects and wait for hydrated counts; confirm the employee sees only intended work.",
        "Open one objective, OKR, NCR, Organization, and Fix-It surface.",
        "Open Create, select a type, and stop before submission unless the rollout script calls for a real record.",
        "Install or launch the PWA on the actual device; enable push only with the employee present.",
        "Receive one email and one push, then confirm the notification opens the intended record.",
    ])

    new_page(doc, "05 / FIRST-LOAD TRUTH", "Believable zeros appear before real data", "This is the most dangerous product-behavior issue because it looks like valid business truth, not a loading state.")
    add_callout(doc, "Measured behavior", "On a 390×844 mobile viewport, NCR displayed zero cards and 0 of 0 records for more than five seconds; the real 74 / 45 / 48 / 296 counts appeared at approximately 7.36 seconds. Dashboard, OKR, KPI, Organization, and Fix-It showed the same false-empty pattern during this audit, generally resolving within 5–10 seconds.", "red")
    add_figure(doc, ASSETS/"hydration-comparison.jpg", "Before-and-after production evidence. The left states are not screenshots taken too early by accident; they are the user-visible truth until hydration completes.", 6.75)
    add_heading(doc, "Why this matters tomorrow", 2)
    add_bullets(doc, [
        "A new user may conclude the app is empty, broken, or that their permissions are wrong.",
        "A leader may read 0 open NCRs or 0 active objectives as actual operating status.",
        "The state can trigger refreshes, duplicate support tickets, and unnecessary sign-outs during onboarding.",
        "Route-by-route settling creates inconsistent screenshots and makes facilitator troubleshooting harder.",
    ])
    add_heading(doc, "Safe temporary operating instruction", 2)
    add_callout(doc, "Until the interface has explicit loading semantics", "Facilitators should wait 10 seconds after every first route open. If data remains at zero after 15 seconds, hard-refresh once, then stop the cohort and escalate. Do not make a management decision from a first-load zero.", "amber")

    new_page(doc, "06 / COMMUNICATION", "The auto-open Daily bulletin is stale on launch eve", "The first message shown to a user contradicts the August 5 rollout schedule.")
    add_figure(doc, EVIDENCE/"desktop/01-entry.jpg", "Observed on entry August 4, 2026: current-day masthead with a rollout story dated Wednesday, June 24, 2026.", 6.85)
    add_heading(doc, "Observed conflict", 2)
    add_table(doc, ["Element", "Observed text", "Risk"], [
        ("Current masthead", "Tuesday, Aug 4, 2026", ("Current", "green")),
        ("Headline", "SandPro OMP goes company-wide next Wednesday", ("Misleading", "red")),
        ("Article date", "Wednesday, June 24, 2026 · 8:34 AM", ("Stale", "red")),
        ("Body", "Moves company-wide next Wednesday, June 24", ("Contradictory", "red")),
    ], widths=[1.3,4.4,1.2], font_size=8)
    add_callout(doc, "Gate", "A stale auto-open rollout announcement is a trust failure at the exact moment new users are deciding whether the system is current. Refresh it, clearly label it as archive, or suppress it before the first broad cohort — then reopen production and prove the corrected state.", "red")

    new_page(doc, "07 / OBJECTIVES + OKR", "The execution system works, but deadline and KPI semantics need care", "The objective portfolio is populated and owned; most work has no due date.")
    add_figure(doc, ASSETS/"objectives.png", "Authoritative production objective status distribution.", 6.8)
    add_metric_cards(doc, [
        ("90 / 92", "No due date", "98% of objectives", "red"),
        ("92 / 92", "Has an owner", "Ownership present", "green"),
        ("3 / 92", "Measured objectives", "Most are simple status", "amber"),
        ("0", "Blocked", "At audit time", "green"),
    ])
    add_heading(doc, "What the settled views show", 2)
    add_bullets(doc, [
        "Dashboard: 82 active — 55 on track, 13 at risk, 14 not started — plus 10 completed.",
        "OKR: August check-in progress states 0 of 90 lines updated; four company OKRs are visible, followed by department lines.",
        "KPI Command Center: 75% operating health, 82 active objectives, 80% NCR closure, 4 of 5 computed KPIs measured.",
        "Objective execution health is 86%, above its 85% target, but the due-readiness KPI is green at 0 because 90 objectives have no due dates.",
        "Manual scorecards have no data; production has 0 KPI definitions and 0 stored KPI datapoints.",
        "The KPI 'OKR Trust' lens reports no key results even though the OKR workspace contains monthly lines; the two surfaces use different semantics or wiring.",
    ])
    add_figure(doc, EVIDENCE/"desktop/24-okr-settled.jpg", "Settled desktop OKR view. Four company OKRs and department lines are present after hydration.", 6.85)

    new_page(doc, "08 / KPI TRUST", "Green can be technically correct and operationally misleading", "The KPI layer is useful once hydrated, but several readings depend on absent source data.")
    add_figure(doc, EVIDENCE/"desktop/25-kpi-settled.jpg", "Settled KPI Command Center: 75% health, 82 active objectives, 80% NCR closure, manual scorecards absent.", 6.85)
    add_table(doc, ["KPI", "Observed", "Interpretation"], [
        ("Operating health", "75%", "One watch KPI and one missing-data KPI; 4 of 5 measured."),
        ("Objective execution", "86%", "58 active customer-visible objectives; 8 at-risk actions in the KPI lens."),
        ("7-day due readiness", "On target · 0", "Not evidence of schedule health when 90/92 objectives have no due date."),
        ("NCR closure", "80% vs 85% target", "296 of 370 closed; 74 remain open."),
        ("OKR Trust", "No data · 0 stale KRs", "Does not reflect the populated OKR monthly-line workspace."),
        ("Manual scorecards", "No data", "No definitions or datapoints stored."),
    ], widths=[1.6,1.5,3.85], font_size=8)
    add_callout(doc, "Keep in mind", "For the first week, present due-readiness and OKR Trust as data-quality indicators, not performance grades. A green zero is not proof that time-bound execution is under control.", "amber")

    new_page(doc, "09 / NCR", "The application can surface the backlog; the organization cannot yet govern it", "This is an operating-model risk more than a rendering problem.")
    add_figure(doc, ASSETS/"ncr.png", "Authoritative database counts, reconciled against settled Tracker, Closeout, and Analytics views.", 6.8)
    add_table(doc, ["Measure", "Current", "Meaning"], [
        ("Status", "72 open · 2 in progress · 296 closed", "74 work-in-scope records remain open."),
        ("Lifecycle", "63 draft · 11 submitted · 296 closed", "Most open NCRs have not moved beyond draft."),
        ("Critical open", ("48", "red"), "65% of the open backlog is critical."),
        ("Past due", ("45", "red"), "Past-due follow-up dates."),
        ("Older than 45 days", ("74", "red"), "Database query; UI Trend Watch displayed 73."),
        ("Open without owner", ("73", "red"), "Only one open NCR has an owner."),
        ("Open without due date", ("29", "amber"), "Follow-up cannot be schedule-managed."),
        ("Missing main department", ("194", "amber"), "Legacy department triage remains unresolved."),
        ("Closeout ready", ("0", "red"), "All 74 open rows need data."),
        ("Supporting rows", "0 actions · 0 attachments · 0 signatures", "Structured closeout evidence tables are empty."),
    ], widths=[1.7,2.2,3.05], font_size=8)
    add_figure(doc, EVIDENCE/"desktop/26-ncr-tracker-settled.jpg", "Settled NCR Tracker. Cards show 74 open, 45 past due, 48 critical open, 296 closed; the default list shows 72 status=open records.", 6.85)

    new_page(doc, "10 / NCR OPERATING PLAN", "Do not turn backlog cleanup into tomorrow's launch task", "Make the backlog visible, owned, and intentionally sequenced.")
    add_heading(doc, "Launch-day minimum", 2)
    add_bullets(doc, [
        "Name one NCR data steward and one executive escalation owner before the first cohort.",
        "Explain the count definitions: the card's 74 includes open plus in-progress; the default table displays 72 status=open records.",
        "Do not promise that all historical NCRs are clean. State that 194 require legacy department triage and 74 need closeout data.",
        "Freeze KPA imports and bulk department mapping during onboarding unless a named data steward is running a rehearsed procedure.",
        "Escalate the 48 critical open NCRs into a separate review cadence; do not bury them in the general rollout checklist.",
    ])
    add_heading(doc, "First 30 days", 2)
    add_table(doc, ["Window", "Action", "Exit evidence"], [
        ("Day 1–2", "Assign triage owner/group to every open NCR", "74/74 have an accountable contact"),
        ("Week 1", "Resolve the 29 missing due dates and reconcile 73 vs 74 age count", "Date rules documented; dashboard reconciled"),
        ("Week 1–2", "Triage 48 critical open NCRs", "Executive-reviewed disposition and next date"),
        ("Week 2–4", "Map 194 legacy departments", "No unresolved legacy mapping in current operating view"),
        ("Month 1", "Populate action, attachment, and signature evidence", "Closeout-ready count begins moving above zero"),
    ], widths=[1.0,3.55,2.4], font_size=8)
    add_figure(doc, EVIDENCE/"desktop/06-ncr-analytics.jpg", "NCR Analytics is a strong production surface. It clearly exposes aging and criticality once loaded.", 6.85)

    new_page(doc, "11 / NOTIFICATIONS", "The infrastructure works for a pilot; company-wide reach does not", "Channel readiness is constrained by preferences, sign-in, and device subscription — not by a visible failure spike.")
    add_figure(doc, ASSETS/"notifications.png", "SandPro personnel notification coverage. 'Enabled' is a preference row; active push subscription is the stronger device-level signal.", 6.8)
    add_table(doc, ["Signal", "Observed", "Interpretation"], [
        ("In-app enabled", "90 / 90 personnel", "Useful only after employees sign in."),
        ("Email enabled", "3 / 90 personnel", "Not a company-wide delivery channel yet."),
        ("Push enabled", "3 / 90 personnel", "Preference exists for three personnel."),
        ("Active push subscription", "2 / 90 personnel", "Only two SandPro personnel have an active device subscription."),
        ("Daily digest frequency", "67 personnel", "Frequency is set, but only three personnel have email enabled."),
        ("Past 24h email", "4 sent · 0 failed", "Positive infrastructure signal for the opted-in cohort."),
        ("Past 24h push", "6 sent · 35 skipped by preference", "No delivery failures in the sampled channel log; reach remains low."),
    ], widths=[1.65,1.65,3.65], font_size=8)
    add_heading(doc, "Do not do tomorrow", 2)
    add_bullets(doc, [
        "Do not send a blanket company-wide message and assume it reached all 90 employees.",
        "Do not treat 'daily digest' as enabled email delivery when the email channel itself is disabled.",
        "Do not enable push on employee devices without the employee present and without verifying the notification click target.",
        "Do not rely on the Release Smoke Admin's empty personal notification feed as proof that the system has no notifications.",
    ])

    new_page(doc, "12 / MOBILE + PWA", "Responsive behavior is credible; installed-device proof is still missing", "Browser viewport evidence is useful, but it is not iPhone or Android installation evidence.")
    add_figure(doc, ASSETS/"mobile-surfaces.jpg", "Observed at 390×844 after hydration. The mobile header, bottom navigation, Create, OKR, NCR, and Organization surfaces reflow coherently.", 6.7)
    add_heading(doc, "What passed", 2)
    add_bullets(doc, [
        "No major horizontal crop was observed at 390×844 across the five core surfaces.",
        "Bottom navigation exposes Work, OKR, NCR, Fix-It, and Org with recognizable labels.",
        "The Create wizard opens and remains usable at mobile width; no record was submitted.",
        "The live manifest declares standalone display, portrait-primary orientation, maskable icons, start URL, and four shortcuts.",
        "The live service worker is reachable, network-first for navigation, bypasses Supabase/API requests, and includes push click handling.",
    ])
    add_heading(doc, "What remains unproven", 2)
    add_bullets(doc, [
        "Actual install, launch, icon, safe-area behavior, offline shell, and update adoption on a physical iPhone or Android device.",
        "A real employee push permission prompt, active device subscription, notification arrival, and deep-link open.",
        "PWA cache/version match after tomorrow's rollout; the shell currently identifies cache version sandpro-omp-shell-v10.",
        "The false-zero window on slower cellular or field conditions; the 7.36-second result came from the in-app browser environment.",
    ])

    new_page(doc, "13 / SECURITY + PLATFORM", "No active incident signal, but several release-hardening gaps remain", "Security advisors are warnings, not proof of compromise. They still deserve explicit ownership before broad access.")
    add_heading(doc, "Security findings", 2)
    sec_rows = [
        (("P0 / ACCEPT OR FIX", "red"), "handle_new_user() is SECURITY DEFINER and executable by anon and authenticated roles", "Advisor-confirmed; direct privilege query confirmed both grants. The function is a trigger function, which limits ordinary RPC use, but the exposed grant is unnecessary attack surface."),
        (("P1", "amber"), "Leaked-password protection is disabled", "Broad rollout increases the value of preventing known-compromised passwords."),
        (("P1", "amber"), "Five functions have mutable search_path", "Harden function search paths to reduce object-resolution risk."),
        (("P1", "amber"), "Observed response lacks CSP and common hardening headers", "HSTS is present; no CSP, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, or frame restriction was observed on the root response."),
        (("P2", "blue"), "172 performance advisories", "76 unindexed foreign keys, 56 auth RLS init-plan warnings, 22 multiple-permissive-policy warnings, 17 unused indexes, and one auth DB connection advisory."),
    ]
    add_table(doc, ["Priority", "Finding", "Why it matters"], sec_rows, widths=[1.25,2.85,2.85], font_size=7.7)
    add_heading(doc, "Positive platform signals", 2)
    add_bullets(doc, [
        "Strict-Transport-Security is present with a two-year max-age.",
        "Relevant public tables inspected have RLS enabled.",
        "Vercel reports the production deployment READY and returned no error/warning log rows for the prior 24 hours.",
        "The live static shell and service worker are cache hits at the edge, reducing asset delivery load.",
    ])
    add_heading(doc, "Advisor references", 2)
    for label, url in [
        ("Supabase function search_path lint", "https://supabase.com/docs/guides/database/database-linter?lint=0011_function_search_path_mutable"),
        ("Anonymous SECURITY DEFINER execution lint", "https://supabase.com/docs/guides/database/database-linter?lint=0028_anon_security_definer_function_executable"),
        ("Authenticated SECURITY DEFINER execution lint", "https://supabase.com/docs/guides/database/database-linter?lint=0029_authenticated_security_definer_function_executable"),
        ("Supabase leaked-password protection", "https://supabase.com/docs/guides/auth/password-security#password-strength-and-leaked-password-protection"),
    ]:
        p = doc.add_paragraph(style="NDAI Bullet")
        set_para(p, after=3)
        add_hyperlink(p, label, url)

    new_page(doc, "14 / ORGANIZATION + ADMIN", "The roster and assignment groups exist; production hygiene needs one pass", "Admin capacity is present, but QA identities are visible inside employee-facing organization surfaces.")
    add_metric_cards(doc, [
        ("97", "People displayed", "Includes 7 support / QA", "lav"),
        ("0", "Org header groups", "Different from assignment groups", "amber"),
        ("10", "Assignment groups", "All active", "green"),
        ("129", "Group memberships", "Across 10 groups", "sand"),
    ])
    add_heading(doc, "Production hygiene findings", 2)
    add_bullets(doc, [
        "Seven support/QA identities are present in the 97-person production roster: Agent Fixit, Andrew Emmel, Mobile Zero-Day QA, Release Smoke Admin, Release Smoke Member, SandPro QA Agent, and Thrawn NDAI.",
        "The mobile org chart presents Agent Fixit as a company root before the CEO and exposes platform/support identities to employees.",
        "The Organization header says 0 groups while the Admin assignment-group panel has 10 active groups. These are different concepts, but the labeling will look contradictory to new administrators.",
        "A completed objective titled 'Test' remains in production; a second objective containing 'test failures' appears to be legitimate operational wording and should not be removed blindly.",
        "All Personnel and Office Personnel are provisioned; the All Personnel group is the appropriate rollout audience only after activation/channel proof is complete.",
    ])
    add_figure(doc, EVIDENCE/"desktop/13-admin-groups.jpg", "Admin assignment groups: 10 active groups, including All Personnel and Office Personnel.", 6.85)

    new_page(doc, "15 / UX + ACCESSIBILITY", "A cohesive light interface with a few high-leverage trust and keyboard gaps", "The visual system is consistent across desktop and mobile; the findings below are about behavior and access, not style preference.")
    add_table(doc, ["Priority", "Finding", "Evidence / impact"], [
        (("P0", "red"), "False-empty loading states", "Zero and 'no data' screens persist for 5–8 seconds before production values appear."),
        (("P0", "red"), "Stale auto-open Daily", "Launch date is June 24 on August 4."),
        (("P1", "amber"), "Clickable cards lack button semantics", "Desktop KPI cards use cursor-pointer divs with no role/tab stop; keyboard users may not reach drilldowns."),
        (("P1", "amber"), "Profile trigger lacks button semantics", "Clickable profile container is a div without role, tabIndex, or aria-expanded."),
        (("P1", "amber"), "Count-definition ambiguity", "Dashboard Past Due objective KPI is 0 while the combined list has 45 past-due NCRs; NCR card shows 74 while default table shows 72."),
        (("P1", "amber"), "Create step indicator starts at 3 of 5", "The first visible question is step 1 while the progress note says originator/assignee are captured; explain or simplify during onboarding."),
        (("P2", "blue"), "Organization group labeling", "0 groups in Organization vs 10 assignment groups in Admin reads as contradictory."),
    ], widths=[.8,2.25,3.95], font_size=8)
    add_heading(doc, "Visual strengths", 2)
    add_bullets(doc, [
        "Clear information hierarchy, restrained SandPro orange, readable cards, and consistent navigation.",
        "Responsive mobile header and bottom navigation preserve the five core surfaces without major crop.",
        "NCR Analytics and Closeout make the quality backlog legible and actionable once hydrated.",
        "Create uses a guided form with re-selectable type choices and clear department/assignment structure.",
        "Admin Users, Groups, Departments, Reports, Export, and Settings form a coherent operational rail.",
    ])
    add_figure(doc, EVIDENCE/"desktop/27-create-wizard.jpg", "Desktop Create wizard opened successfully; no production record was submitted.", 6.85)

    new_page(doc, "16 / ROLLOUT RUNBOOK", "A staffed, observable launch plan for August 5", "Use this as the operating checklist, not as permission to skip the hard gates.")
    runbook_rows = [
        ("T−24 to T−2h", "Launch commander", "Confirm current deploy ID, fresh Daily, cohort roster, ordinary-user proof, support channel, and stop criteria."),
        ("T−2h", "Activation lead", "Verify first-wave credentials; run one real-user sign-in and password recovery if applicable."),
        ("T−90m", "Platform owner", "Open Dashboard, OKR, KPI, NCR, Org, Fix-It; wait for hydration; record settled counts."),
        ("T−60m", "Comms owner", "Receive a real employee email; inspect inbox and junk; verify sender/subject/links."),
        ("T−45m", "Device owner", "Install/launch PWA on employee device; enable push; receive and open one push."),
        ("T0", "Facilitator", "Activate 5–10 employees. Every person completes sign-in, navigation, and notification checklist."),
        ("T+15m", "Launch commander", "Hold/advance decision. Do not start the next cohort until all current users are green."),
        ("T+60m", "Data steward", "Check new records, ownership visibility, notifications, Fix-It, and unexpected permission gaps."),
        ("EOD", "Executive owner", "Publish adoption totals, incidents, deferred issues, and next-day cohort decision."),
    ]
    add_table(doc, ["When", "Owner", "Action / proof"], runbook_rows, widths=[1.15,1.5,4.35], font_size=8)
    add_heading(doc, "Stop-the-line criteria", 2)
    add_bullets(doc, [
        "More than 10% of the current cohort cannot sign in or recover access within 10 minutes.",
        "Any ordinary employee sees administrator-only controls or data outside intended visibility.",
        "A route remains falsely empty after 15 seconds and one hard refresh.",
        "A notification opens the wrong record, fails on the employee device, or is sent to the wrong audience.",
        "Repeated 5xx errors, a visible crash, or data writes that do not persist after reread.",
    ])
    add_callout(doc, "Rollback posture", "Stop adding cohorts. Keep the successful pilot group active. Preserve data, logs, and screenshots. Roll back a deployment only if a specific deployment defect is confirmed and the previous production artifact is known-good; adoption trouble alone is not a reason for a database rollback.", "blue")

    new_page(doc, "17 / ACCEPTANCE CHECKLIST", "Go / no-go decision sheet", "Every red item needs evidence or an accountable risk acceptance. A verbal 'should work' is not evidence.")
    checklist_rows = [
        ("□", "Daily bulletin reflects August 5 rollout or is intentionally suppressed", "Comms owner", "Screenshot + timestamp"),
        ("□", "First-wave employee accounts sign in on their actual devices", "Activation lead", "Named cohort list"),
        ("□", "Ordinary employee sees intended records and no admin-only surfaces", "Product owner", "Screen recording / screenshots"),
        ("□", "False-zero behavior mitigated or facilitator protocol accepted", "Platform owner", "Loading proof or signed playbook"),
        ("□", "Email reaches inbox or documented junk path", "Comms owner", "Mailbox screenshot / trace"),
        ("□", "Push arrives and opens correct record on a real device", "Device owner", "Device video / screenshot"),
        ("□", "SECURITY DEFINER grants and leaked-password setting fixed or accepted", "Security owner", "Advisor rerun / risk signoff"),
        ("□", "NCR triage and escalation owners named", "Quality owner", "Owner list + cadence"),
        ("□", "Production deploy ID and rollback artifact recorded", "Launch commander", "Deployment record"),
        ("□", "Stop criteria and support channel communicated", "Launch commander", "Launch message"),
    ]
    add_table(doc, ["", "Gate", "Accountable owner", "Required evidence"], checklist_rows, widths=[.35,3.15,1.45,2.05], font_size=8)
    add_heading(doc, "Decision record", 2)
    add_table(doc, ["Decision", "Owner", "Time", "Conditions / accepted risk"], [
        ("GO / HOLD", "________________", "____________", "________________________________________________"),
        ("Next cohort", "________________", "____________", "________________________________________________"),
        ("EOD status", "________________", "____________", "________________________________________________"),
    ], widths=[1.15,1.55,1.15,3.15], font_size=8.5)
    add_callout(doc, "Final recommendation", "Hold the unguarded company-wide launch. If the hard gates are cleared, proceed with a controlled, staffed rollout in small waves and treat activation, first-load truth, and channel proof as the first day's primary success measures.", "red")

    new_page(doc, "APPENDIX A", "Evidence index and reconciliation notes", "The report embeds representative screenshots; the full evidence set remains alongside the source builder in the audit folder.")
    add_table(doc, ["Evidence", "Observed state", "Use in decision"], [
        ("01-entry.jpg", "Stale auto-open Daily", "Communication gate"),
        ("02-dashboard.jpg", "Settled desktop dashboard", "Objective and mixed-list snapshot"),
        ("24-okr-settled.jpg", "Populated OKR workspace", "OKR availability / 90-line check-in"),
        ("25-kpi-settled.jpg", "Computed KPI view", "Performance and semantic review"),
        ("26-ncr-tracker-settled.jpg", "Populated NCR tracker", "Backlog counts / count definitions"),
        ("05-ncr-closeout.jpg", "0 ready, 74 need data", "Closeout readiness"),
        ("06-ncr-analytics.jpg", "Trend Watch and aging", "Quality operating risk"),
        ("08-ncr-triage.jpg", "194 legacy mappings", "Data stewardship"),
        ("13-admin-groups.jpg", "10 active assignment groups", "Audience readiness"),
        ("22-fixit-settled.jpg", "0 active, 28 archived", "Support readiness"),
        ("27-create-wizard.jpg", "Create opened; no submit", "Core workflow availability"),
        ("mobile/02,03,04,07,06", "Responsive core surfaces", "Mobile-browser readiness"),
    ], widths=[1.65,2.55,2.75], font_size=7.7)
    add_heading(doc, "Reconciliation notes", 2)
    add_bullets(doc, [
        "NCR 74 vs 72: the high-level card counts open plus in-progress; the default selected table shows 72 status=open rows.",
        "NCR age 73 vs 74: production Analytics displayed 73 older than 45 days; the read-only database query at 08:40 CDT returned all 74 open rows older than 45 days. Reconcile the date boundary and source logic.",
        "OKR 90 vs objectives 92: the OKR progress bar tracks 90 open lines; the objective table contains 92 total rows including 10 completed. All 92 have OKR linkage fields populated in the production data model.",
        "Organization 0 groups vs Admin 10 groups: the Organization header appears to count org-chart groups; Admin counts assignment groups. Label the concepts explicitly.",
        "Past Due 0 vs 45: dashboard KPI describes objective deadlines; the combined list filter includes NCR follow-up deadlines. The UI needs clearer scope labels.",
    ])

    new_page(doc, "APPENDIX B", "Audit limitations and data handling", "A comprehensive audit is strongest when its exclusions are explicit.")
    add_heading(doc, "Not tested", 2)
    add_bullets(doc, [
        "Destructive or write workflows: create submission, objective/NCR updates, owner assignment, org edits, imports, exports, user changes, password changes, group changes, Fix-It posting, commenting, or archival.",
        "A real SandPro employee session or ordinary-role credential; the visual audit used Release Smoke Admin.",
        "Physical iPhone, iPad, or Android installation; offline behavior; background notification delivery; cellular performance.",
        "Microsoft 365 tenant logs, inbox/junk placement, sender-domain authentication, or employee mailbox delivery.",
        "Staging mutations, local release preflight, schema migration, build reproducibility, or deployment rollback execution.",
        "Full API/Postgres log export: the connected log endpoint did not return those services; Vercel and sampled Auth logs were used instead.",
    ])
    add_heading(doc, "Data handling", 2)
    add_bullets(doc, [
        "Read-only SQL was used for aggregate metrics and named support/activation verification.",
        "Screenshots contain production names and business records; this report should remain internal to SandPro/NDAI until redacted.",
        "No production secret, password, token, or service-role credential is included in the deliverable.",
        "The working repository was already dirty; unrelated changes were preserved and no deploy was attempted.",
    ])
    add_callout(doc, "Evidence standard", "A visible screen, passing deployment status, table count, or absence of logs is never enough by itself. The rollout gates require the behavior to work for the intended ordinary employee, on the intended device, with the intended data and communication channel.", "lav")

    # Core properties
    props = doc.core_properties
    props.title = "SandPro OMP Final Pre-Rollout Audit"
    props.subject = "Read-only production readiness audit for objectivetracker.net"
    props.author = "NDAI / Codex"
    props.keywords = "SandPro OMP, production audit, rollout readiness, objectivetracker.net, NDAI"
    props.comments = "Generated from read-only production evidence captured August 4, 2026."
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
