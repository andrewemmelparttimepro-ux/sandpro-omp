from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("SandPro_OMP_Quick_Update_for_Merci.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DEE8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_section_heading(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 127, 2)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(9.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(2)
run = title.add_run("SandPro OMP: Quick Update for Merci")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(21)
run.font.color.rgb = RGBColor(17, 24, 39)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
run = subtitle.add_run("Plain-English summary of the latest cleanup and testing focus")
run.font.name = "Calibri"
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(107, 114, 128)

intro_table = doc.add_table(rows=1, cols=1)
intro_table.autofit = False
intro_table.columns[0].width = Inches(7.0)
intro_cell = intro_table.cell(0, 0)
set_cell_shading(intro_cell, "FFF4EC")
set_cell_border(intro_cell, "FDBA74")
intro = intro_cell.paragraphs[0]
intro.paragraph_format.space_after = Pt(0)
intro.paragraph_format.line_spacing = 1.08
intro_run = intro.add_run(
    "Hi Merci - here is the quick version of what was tightened up so you can help test the platform with confidence. "
    "The goal is simple: make the app easier to understand, easier to maintain, and less likely to surprise you during real SandPro work."
)
intro_run.font.name = "Calibri"
intro_run.font.size = Pt(9.5)
intro_run.font.color.rgb = RGBColor(31, 41, 55)

add_section_heading(doc, "What changed")
for item in [
    "You can now help maintain the Organization page. You should be able to update names, titles, departments, and who someone reports to.",
    "Platform role changes are still protected. Day-to-day org cleanup is available to you, but deeper access changes stay limited to executive admins.",
    "Helpful in-app tips were added for new workflows. Each tip can be closed, and then reopened later from a small question-mark help button.",
    "Objective tagging and workflow guidance is clearer, including when to tag a teammate, when to use the workflow tracker, and when to mention someone in a message.",
    "File, Daily Brief, and notification areas now include short guidance so the app explains what to do without needing a separate training call.",
    "Additional validation checks were added so permissions, help prompts, login recovery, and release basics are tested before changes are handed off.",
]:
    add_bullet(doc, item)

add_section_heading(doc, "What this means for you")
for item in [
    "If Jake or Andrew asks for org chart cleanup, you should be able to do that directly from the Organization page.",
    "When a feature is new or unfamiliar, look for the short tip card or the small question-mark button.",
    "If something does not work the way the tip says it should, treat that as useful beta feedback and send it over right away.",
]:
    add_bullet(doc, item)

add_section_heading(doc, "What is coming next")
roadmap_table = doc.add_table(rows=1, cols=1)
roadmap_table.autofit = False
roadmap_table.columns[0].width = Inches(7.0)
roadmap_cell = roadmap_table.cell(0, 0)
set_cell_shading(roadmap_cell, "F8FAFC")
set_cell_border(roadmap_cell, "CBD5E1")
roadmap = roadmap_cell.paragraphs[0]
roadmap.paragraph_format.space_after = Pt(0)
roadmap.paragraph_format.line_spacing = 1.08
roadmap_run = roadmap.add_run(
    "A longer roadmap is being prepared for Jake, Andrew, and you. It will lay out the next phase in plain English: "
    "cleaner notifications, better mobile use, the workflow tracker, file/export improvements, org chart cleanup, and the path toward a wider SandPro team rollout."
)
roadmap_run.font.name = "Calibri"
roadmap_run.font.size = Pt(9.5)
roadmap_run.font.color.rgb = RGBColor(31, 41, 55)

footer = doc.add_paragraph()
footer.paragraph_format.space_before = Pt(8)
footer.paragraph_format.space_after = Pt(0)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Prepared for internal SandPro OMP beta coordination.")
footer_run.font.name = "Calibri"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = RGBColor(107, 114, 128)

doc.save(OUT)
print(OUT)
