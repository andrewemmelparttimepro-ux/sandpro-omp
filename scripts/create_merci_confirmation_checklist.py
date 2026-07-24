from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/andrewemmel/Documents/New project/sandpro-omp")
OUTPUT = ROOT / "docs" / "Merci_OMP_Confirmation_Checklist_2026-07-23.docx"
LOGO = ROOT / "public" / "brand" / "sandpro-omp-logo.png"

NAVY = "17324D"
ORANGE = "F47C20"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "25364A"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_ORANGE = "FFF4E8"
BORDER = "CDD5DF"
WHITE = "FFFFFF"

GROUPS = [
    ("Dispatch", "Dustin Saunders; Gershom Dingal; Luke Feil; Shawn Cockrell"),
    ("Field Service Managers", "Isaac Badillo; Zedek Harris"),
    ("Trainers", "Bryce Christoffersen; Brad Beck"),
    ("Sales Team", "John Sommerfeld; Jon Ostby; Brandon Schatz; Josh Pfeifer; Larry Debold"),
    ("CP Shop Leads", "Kelby Kraft; Eric Macy"),
    ("Flowback Shop Leads", "Matthew Bornschein"),
    ("Wellhead Shop Leads", "Thomas Goldsberry; Jeramiah Walls"),
    ("Leadership / Business Team", "Jake Feil; Joshua Blackaby; Andrew Emmel"),
]

GROUP_QUESTIONS = [
    "Should Kelby Kraft and Eric Macy both be in CP Shop Leads?",
    "Is Matthew Bornschein the only Flowback Shop Lead?",
    "Should Thomas Goldsberry and Jeramiah Walls both be in Wellhead Shop Leads?",
    "Is the Sales Team list complete for shared goals?",
    "Should Leadership / Business Team include anyone beyond Jake, Joshua, and Andrew?",
]

IDENTITIES = [
    ("Mark Elliott", "melliott@sandpro.com"),
    ("Bryce Christoffersen", "bchristoffersen@sandpro.com"),
    ("Brad Beck", "bbeck@sandpro.com"),
    ("Derek Holen", "dholen@sandpro.com"),
    ("Hunter Haataia", "hhaataia@sandpro.com"),
    ("Eric Macy", "emacy@sandpro.com"),
    ("Adrian Blackaby", "ablackaby@sandpro.com"),
    ("Jonathan Shelstad", "jshelstad@sandpro.com"),
    ("Carter Caswell", "ccaswell@sandpro.com"),
    ("Edwin Blanco", "eblanco@sandpro.com"),
    ("Marvin Roa Baca", "mroabaca@sandpro.com"),
    ("Aidyn Ross", "aross@sandpro.com"),
    ("Reene Calderon", "rcalderon@sandpro.com"),
    ("Garl McGraw", "gmcgraw@sandpro.com"),
    ("Julius Williams", "jwilliams@sandpro.com"),
    ("Hunter Jones", "hjones@sandpro.com"),
    ("Richard Griffin", "rgriffin@sandpro.com"),
    ("Phillip Leviner", "pleviner@sandpro.com"),
    ("Bob Young", "byoung@sandpro.com"),
    ("Curtis Jones", "cjones@sandpro.com"),
    ("Corey Sharkey", "csharkey@sandpro.com"),
    ("Logan Howard", "lhoward@sandpro.com"),
    ("Austin Dees", "adees@sandpro.com"),
    ("Marcos Vega", "mvega@sandpro.com"),
    ("Brian Brower", "bbrower@sandpro.com"),
    ("Kris Trone", "ktrone@sandpro.com"),
    ("Fred Floyd Jr.", "ffloydjr@sandpro.com"),
    ("Joseph Dingal", "jdingal@sandpro.com"),
    ("Luke Feil", "lfeil@sandpro.com"),
    ("Shane Vogel", "svogel@sandpro.com"),
    ("Abel Lua", "alua@sandpro.com"),
    ("Shawn Cockrell", "scockrell@sandpro.com"),
    ("Bill Anderson", "banderson@sandpro.com"),
    ("Jake Beck", "jbeck@sandpro.com"),
    ("Able Conley", "aconley@sandpro.com"),
    ("Wyatt Phipps", "wphipps@sandpro.com"),
    ("Austin Griffin", "agriffin@sandpro.com"),
    ("Jean Bazile", "jbazile@sandpro.com"),
    ("Kobie Jones", "kjones@sandpro.com"),
    ("Josef Mcconnell", "jmcconnell@sandpro.com"),
    ("Kevin Johnson", "kjohnson@sandpro.com"),
    ("Jeremy Tate", "jtate@sandpro.com"),
    ("Jerimiah Howard", "jhoward@sandpro.com"),
    ("Nick Reiter", "nreiter@sandpro.com"),
    ("Dexter Sotelo", "dsotelo@sandpro.com"),
    ("Dion Carter", "dcarter@sandpro.com"),
    ("Josh Preston", "jpreston@sandpro.com"),
    ("Ben Blanco", "bblanco@sandpro.com"),
    ("Ryan Yensal", "ryensal@sandpro.com"),
    ("Jerry Hassler", "jhassler@sandpro.com"),
    ("Matthew Bornschein", "mbornschein@sandpro.com"),
    ("Alex Mora", "amora@sandpro.com"),
    ("Dylan DeBold", "ddebold@sandpro.com"),
    ("Kaden Mackay", "kmackay@sandpro.com"),
    ("Travis Bell", "tbell@sandpro.com"),
    ("Andrii Navrotskyi", "anavrotskyi@sandpro.com"),
    ("Oleg Postolatii", "opostolatii@sandpro.com"),
    ("Jeramiah Walls", "jwalls@sandpro.com"),
    ("Brandon Besselman", "bbesselman@sandpro.com"),
    ("Zachary Campana", "zcampana@sandpro.com"),
    ("Wyatt Decoteau", "wdecoteau@sandpro.com"),
    ("Jovan Blanco", "jblanco@sandpro.com"),
    ("Tobay Hall", "thall@sandpro.com"),
    ("Keion Louden", "klouden@sandpro.com"),
]

DUAL_MANAGER_EMPLOYEES = [
    "Garl McGraw",
    "Julius Williams",
    "Hunter Jones",
    "Richard Griffin",
    "Phillip Leviner",
    "Bob Young",
    "Curtis Jones",
    "Corey Sharkey",
    "Logan Howard",
    "Austin Dees",
    "Marcos Vega",
    "Brian Brower",
    "Kris Trone",
    "Fred Floyd Jr.",
    "Joseph Dingal",
    "Shane Vogel",
    "Abel Lua",
    "Bill Anderson",
    "Jake Beck",
    "Able Conley",
    "Wyatt Phipps",
    "Austin Griffin",
    "Jean Bazile",
    "Kobie Jones",
    "Josef Mcconnell",
    "Kevin Johnson",
    "Jeremy Tate",
    "Jerimiah Howard",
    "Nick Reiter",
    "Dexter Sotelo",
    "Dion Carter",
]


def set_run_font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_cell_text(cell, size=9.2, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def style_header_row(row):
    set_repeat_table_header(row)
    for cell in row.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        style_cell_text(cell, size=9, bold=True, color=NAVY)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_running_furniture(section):
    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_number(fp)


def add_title_block(doc):
    logo_p = doc.add_paragraph()
    logo_p.paragraph_format.space_before = Pt(2)
    logo_p.paragraph_format.space_after = Pt(12)
    logo_shape = logo_p.add_run().add_picture(str(LOGO), width=Inches(1.7))
    logo_shape._inline.docPr.set("descr", "SandPro OMP logo")
    logo_shape._inline.docPr.set("title", "SandPro OMP")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    kr = kicker.add_run("ROSTER AND REPORTING REVIEW")
    set_run_font(kr, size=9.5, bold=True, color=ORANGE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("Items for Merci to Confirm")
    set_run_font(tr, size=28, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run(
        "A plain-language review of rotating groups, proposed SandPro email identities, "
        "manager reporting lines, and one existing account."
    )
    set_run_font(sr, size=12.5, color=MUTED)

    metadata = doc.add_table(rows=2, cols=2)
    set_table_geometry(metadata, [4680, 4680], indent_dxa=120)
    set_table_borders(metadata, color="E5E7EB", size="4")
    values = [
        ("Prepared for", "Mercileidy Jimenez"),
        ("Review date", "July 23, 2026"),
        ("Status", "Confirmation requested"),
        ("How to respond", "Mark the form or reply with corrections"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = metadata.rows[idx // 2].cells[idx % 2]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        lr = p.add_run(label.upper() + "\n")
        set_run_font(lr, size=8, bold=True, color=MUTED)
        vr = p.add_run(value)
        set_run_font(vr, size=10, bold=True, color=INK)
        set_cell_shading(cell, LIGHT_GRAY if idx % 2 == 0 else WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text, fill=PALE_ORANGE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color=ORANGE, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    lr = p.add_run(label + "  ")
    set_run_font(lr, size=10, bold=True, color=NAVY)
    tr = p.add_run(text)
    set_run_font(tr, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_section_heading(doc, title, subtitle=None):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(title)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(8)
        sp.paragraph_format.keep_with_next = True
        r = sp.add_run(subtitle)
        set_run_font(r, size=10, color=MUTED)


def add_summary_table(doc):
    add_section_heading(doc, "What needs a decision")
    rows = [
        ("Rotating groups", "Confirm eight optional, editable group rosters."),
        ("Login identities", "Confirm 64 proposed @sandpro.com addresses before notifications are enabled."),
        ("Reporting lines", "Confirm whether Isaac remains the primary manager for 31 dual-manager rows."),
        ("Existing account", "Choose whether Keith Mappes remains active, is archived, or returns to the roster."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Review area"
    table.rows[0].cells[1].text = "Decision needed"
    for label, decision in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = decision
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    style_header_row(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
        style_cell_text(row.cells[0], size=9.4, bold=True, color=NAVY)
        style_cell_text(row.cells[1], size=9.4)


def add_groups_section(doc):
    add_section_heading(
        doc,
        "Rotating group membership",
        "Groups are optional and editable. They provide shared visibility when responsibility rotates; "
        "they do not replace the single accountable owner on an individual task.",
    )
    table = doc.add_table(rows=1, cols=3)
    headers = ["Group", "Current members", "OK"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for group, members in GROUPS:
        cells = table.add_row().cells
        cells[0].text = group
        cells[1].text = members
        cells[2].text = "[   ]"
    set_table_geometry(table, [2200, 6360, 800])
    set_table_borders(table)
    style_header_row(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
        style_cell_text(row.cells[0], size=9.2, bold=True, color=NAVY)
        style_cell_text(row.cells[1], size=9.2)
        style_cell_text(row.cells[2], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    q = doc.add_paragraph(style="Heading 2")
    q.paragraph_format.keep_with_next = True
    q.add_run("Specific questions")
    questions = doc.add_table(rows=1, cols=2)
    questions.rows[0].cells[0].text = "Please confirm"
    questions.rows[0].cells[1].text = "Answer / correction"
    for question in GROUP_QUESTIONS:
        cells = questions.add_row().cells
        cells[0].text = question
        cells[1].text = ""
    set_table_geometry(questions, [6800, 2560])
    set_table_borders(questions)
    style_header_row(questions.rows[0])
    for row in questions.rows[1:]:
        prevent_row_split(row)
        style_cell_text(row.cells[0], size=9.2)
        style_cell_text(row.cells[1], size=9.2)


def add_identities_section(doc):
    add_section_heading(
        doc,
        "Proposed SandPro login identities",
        "These 64 addresses follow the established first-initial-plus-surname pattern. No invitation "
        "emails were sent, and email notifications remain disabled until each address is confirmed.",
    )
    add_callout(
        doc,
        "Fast approval:",
        "If every address is correct, write “All 64 approved” here: _______________________________",
        fill="EEF6FF",
    )

    add_identity_table(doc, IDENTITIES[:8])
    doc.add_page_break()
    add_identity_table(doc, IDENTITIES[8:36])
    doc.add_page_break()
    add_identity_table(doc, IDENTITIES[36:])


def add_identity_table(doc, identities):
    table = doc.add_table(rows=1, cols=4)
    headers = ["Employee", "Proposed email address", "OK", "Correction"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for name, email in identities:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = email
        cells[2].text = "[   ]"
        cells[3].text = ""
    set_table_geometry(table, [2700, 3860, 700, 2100])
    set_table_borders(table, color="D7DEE7", size="5")
    style_header_row(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
        style_cell_text(row.cells[0], size=8.7, bold=True, color=NAVY)
        style_cell_text(row.cells[1], size=8.7)
        style_cell_text(row.cells[2], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell_text(row.cells[3], size=8.7)


def add_reporting_section(doc):
    doc.add_page_break()
    add_section_heading(
        doc,
        "Dual-manager reporting rows",
        "The workbook lists both Isaac Badillo and Zedek Harris for the employees below. OMP uses one "
        "primary org-chart manager, so Isaac is currently shown as primary. Both remain members of the "
        "Field Service Managers rotating group.",
    )
    add_callout(
        doc,
        "Decision:",
        "Keep Isaac as primary for every person, or mark the specific rows that should use Zedek.",
        fill="EEF6FF",
    )

    add_reporting_table(doc, DUAL_MANAGER_EMPLOYEES[:18])
    doc.add_page_break()
    add_reporting_table(doc, DUAL_MANAGER_EMPLOYEES[18:])


def add_reporting_table(doc, employees):
    table = doc.add_table(rows=1, cols=4)
    headers = ["Employee", "Current primary", "Use Zedek?", "Notes / correction"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for name in employees:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = "Isaac Badillo"
        cells[2].text = "[   ]"
        cells[3].text = ""
    set_table_geometry(table, [3000, 2100, 1500, 2760])
    set_table_borders(table, color="D7DEE7", size="5")
    style_header_row(table.rows[0])
    for row in table.rows[1:]:
        prevent_row_split(row)
        style_cell_text(row.cells[0], size=8.8, bold=True, color=NAVY)
        style_cell_text(row.cells[1], size=8.8)
        style_cell_text(row.cells[2], size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        style_cell_text(row.cells[3], size=8.8)


def add_existing_account_section(doc):
    add_section_heading(
        doc,
        "Existing account not listed in the new workbook",
        "This account was not deleted or disabled while the roster was updated.",
    )
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Account"
    table.rows[0].cells[1].text = "Choose one"
    table.rows[1].cells[0].text = "Keith Mappes\nkmappes@sandpro.com\nI&E Technician"
    table.rows[1].cells[1].text = "[   ] Keep active\n[   ] Archive\n[   ] Add back to the roster"
    set_table_geometry(table, [5200, 4160])
    set_table_borders(table)
    style_header_row(table.rows[0])
    style_cell_text(table.rows[1].cells[0], size=9.4, bold=True, color=NAVY)
    style_cell_text(table.rows[1].cells[1], size=9.4)

    add_section_heading(doc, "Approval")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(
        "If the document is correct as shown, reply “Approved as shown.” Otherwise, mark the changes "
        "directly in this file or return a list of corrections."
    )
    set_run_font(r, size=10.5, color=INK)
    signature = doc.add_table(rows=2, cols=2)
    signature.rows[0].cells[0].text = "Reviewed by: __________________________________"
    signature.rows[0].cells[1].text = "Date: __________________"
    signature.rows[1].cells[0].text = "Final notes:"
    signature.rows[1].cells[1].text = ""
    set_table_geometry(signature, [6200, 3160])
    set_table_borders(signature, color="D7DEE7", size="4")
    for row in signature.rows:
        for cell in row.cells:
            style_cell_text(cell, size=9.4)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.0


def build_document():
    assert len(IDENTITIES) == 64
    assert len(DUAL_MANAGER_EMPLOYEES) == 31
    assert sum(len(members.split("; ")) for _, members in GROUPS) == 21

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_running_furniture(section)
    configure_styles(doc)

    props = doc.core_properties
    props.title = "Items for Merci to Confirm - SandPro OMP"
    props.subject = "Roster, rotating groups, login identities, and reporting-line confirmation"
    props.author = "SandPro OMP"
    props.keywords = "SandPro, OMP, roster, groups, identities, reporting"
    props.comments = "Prepared for Mercileidy Jimenez on July 23, 2026."

    add_title_block(doc)
    add_callout(
        doc,
        "Important:",
        "Nothing in this checklist prevents in-app use. Newly inferred email identities remain "
        "notification-disabled until confirmed.",
    )
    add_summary_table(doc)
    add_groups_section(doc)
    add_identities_section(doc)
    add_reporting_section(doc)
    add_existing_account_section(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
